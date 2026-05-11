"""
doc/USAGE.docx 생성 스크립트.

사용:
    python doc/_build_usage_docx.py

이 스크립트는 .docx 문서를 처음부터 다시 만든다.
USAGE.docx 가 깨졌거나 내용 수정이 필요할 때만 실행하면 된다.
런타임 의존성 아님 — 로컬에서 1회성으로만 돌린다.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(HERE, "USAGE.docx")


# ──────────────────────────────────────────────────────────────
# 스타일 헬퍼
# ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Apple SD Gothic Neo"
        run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Apple SD Gothic Neo"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Apple SD Gothic Neo"
    run.font.size = Pt(11)
    return p


def add_code(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(code_text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # 회색 음영
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    p_pr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # 헤더
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Apple SD Gothic Neo"
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(hdr[i], "1F365D")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 본문
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Apple SD Gothic Neo"
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if r_idx % 2 == 1:
            for c in cells:
                set_cell_bg(c, "F7F9FC")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ──────────────────────────────────────────────────────────────
# 문서 본문
# ──────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "Apple SD Gothic Neo"
    style.font.size = Pt(11)

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ── 표지 ───────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("news-collector 사용 가이드")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x1F, 0x36, 0x5D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run("RSS 수집 · Gemini AI 분석 · Google Sheets / Supabase 저장")
    sub_run.italic = True
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x55, 0x66, 0x88)

    add_divider(doc)
    add_para(
        doc,
        "이 문서는 처음 셋업부터 자동화·트러블슈팅까지의 전체 흐름을 한 권에 정리한 운영 가이드입니다. "
        "GitHub Actions 가 매일 한국 시각 09:07 에 자동으로 RSS 를 수집해 AI 분석을 거쳐 "
        "Google Sheets 와 Supabase 두 곳에 저장합니다.",
    )

    # ── 1. 데이터 흐름 ────────────────────────────────────
    add_heading(doc, "1. 데이터 흐름 한눈에 보기", 1)
    add_code(
        doc,
        "GitHub Actions (매일 KST 09:07)\n"
        "        ↓\n"
        "   RSS 수집  (config/categories.yaml)\n"
        "        ↓\n"
        "   본문 크롤링  (crawl_body: true 인 소스)\n"
        "        ↓\n"
        "   Gemini AI 분석 — 요약 · 감성 · 중요도  (배치 20개)\n"
        "        ↓\n"
        "   ├─►  Google Sheets  (연단위 파일 · 월별 탭)   ← TubeAI 대시보드\n"
        "   └─►  Supabase       (365일 retention)         ← 분석·통계용",
    )
    add_para(
        doc,
        "데이터 소비처는 YoutubeProgram/TubeAI_app 의 “RSS 뉴스 피드” 모듈입니다.",
        italic=True,
    )

    # ── 2. 사전 준비물 ────────────────────────────────────
    add_heading(doc, "2. 사전 준비물", 1)
    add_table(
        doc,
        ["항목", "발급처", "용도"],
        [
            ["Google OAuth Client", "Google Cloud Console → Credentials", "Sheets / Drive 접근"],
            ["Gemini API Key", "https://aistudio.google.com/apikey", "AI 요약·감성·중요도 분석"],
            ["Google Drive 폴더", "drive.google.com (폴더 URL 의 ID)", "결과 시트 저장 위치"],
            ["Supabase 프로젝트", "https://supabase.com (선택)", "365일 retention 저장소"],
            ["Python 3.11+", "https://www.python.org", "로컬 실행 환경"],
        ],
    )
    add_para(
        doc,
        "Gemini 와 Supabase 는 선택입니다. 키가 없으면 각각 규칙 기반 fallback / Sheets 단독 저장으로 동작합니다.",
        italic=True,
    )

    # ── 3. 최초 셋업 ─────────────────────────────────────
    add_heading(doc, "3. 최초 셋업 (로컬)", 1)

    add_heading(doc, "3-1. 저장소 클론 & 의존성 설치", 2)
    add_code(
        doc,
        "git clone <this-repo>\n"
        "cd news-collector\n"
        "python -m venv .venv && source .venv/bin/activate\n"
        "pip install -r requirements.txt",
    )

    add_heading(doc, "3-2. .env 파일 만들기", 2)
    add_code(doc, "cp .env.example .env\n# 에디터로 .env 열어서 값 채우기")
    add_para(
        doc,
        ".env 는 .gitignore 로 제외돼 있습니다. 절대 커밋하지 마세요.",
        bold=True,
    )

    add_heading(doc, "3-3. Google OAuth 토큰 발급 (1회)", 2)
    add_para(doc, ".env 에 먼저 client id/secret 만 넣습니다:")
    add_code(doc, "GOOGLE_OAUTH_CLIENT_ID=...\nGOOGLE_OAUTH_CLIENT_SECRET=...")
    add_para(doc, "토큰 발급 스크립트 실행:")
    add_code(doc, "pip install google-auth-oauthlib python-dotenv\npython get_token.py")
    add_para(
        doc,
        "브라우저가 열리고 동의하면 콘솔에 GOOGLE_OAUTH_TOKEN=... 한 줄 JSON 이 출력됩니다. "
        "그 한 줄을 .env 에 복사하면 끝.",
    )
    add_para(doc, "필요 OAuth 스코프:", bold=True)
    add_bullet(doc, "https://www.googleapis.com/auth/spreadsheets")
    add_bullet(doc, "https://www.googleapis.com/auth/drive")

    add_heading(doc, "3-4. 환경변수 요약", 2)
    add_table(
        doc,
        ["변수", "필수", "설명"],
        [
            ["GOOGLE_OAUTH_TOKEN", "✅", "get_token.py 출력값 (JSON 한 줄)"],
            ["GDRIVE_FOLDER_ID", "⚠️", "결과 시트가 저장될 Drive 폴더 ID"],
            ["GEMINI_API_KEY", "⚠️", "없으면 fallback (단순 추출)"],
            ["GDRIVE_SPREADSHEET_ID", "⬜", "특정 시트에 강제 저장하고 싶을 때"],
            ["SUPABASE_URL", "⬜", "Supabase 동시 저장 (선택)"],
            ["SUPABASE_SERVICE_KEY", "⬜", "service_role 키 — 절대 커밋 금지"],
            ["SUPABASE_RETENTION_DAYS", "⬜", "기본 90, Actions 는 365 사용"],
            ["SKIP_GEMINI", "⬜", "1 설정 시 AI 호출 전부 스킵"],
            ["AI_BATCH_SIZE", "⬜", "한 호출에 묶을 기사 수 (기본 20)"],
            ["AI_CALL_DELAY", "⬜", "호출 간 딜레이 초 (기본 1.5)"],
        ],
    )
    add_para(doc, "⚠️ = 필수지만 fallback 있음 · ⬜ = 선택", italic=True)

    # ── 4. 실행 방법 ─────────────────────────────────────
    add_heading(doc, "4. 실행 방법", 1)

    add_heading(doc, "4-1. 로컬 수동 실행", 2)
    add_code(doc, "python src/collector.py config/categories.yaml")
    add_para(doc, "처음 실행하면:")
    add_bullet(doc, "지정한 Drive 폴더에 news-{년도}.xlsx 형태로 시트가 자동 생성됨")
    add_bullet(doc, "월 이름(2026-05) 탭이 생성되고 헤더 행이 작성됨")
    add_bullet(doc, "각 카테고리/소스를 순회하며 RSS → 본문 크롤 → AI 분석 → 행 추가")
    add_bullet(doc, "Supabase 키가 있으면 동시에 Supabase 에도 upsert")

    add_heading(doc, "4-2. Supabase 백필 (선택, 1회)", 2)
    add_para(doc, "기존 시트의 과거 기사를 Supabase 로 한 번에 옮길 때:")
    add_code(
        doc,
        "python scripts/backfill_to_supabase.py 90    # 최근 90일\n"
        "python scripts/backfill_to_supabase.py       # 기본 90일",
    )
    add_para(doc, "url_hash unique 제약으로 중복은 자동 스킵됩니다.")

    add_heading(doc, "4-3. GitHub Actions 자동화", 2)
    add_para(
        doc,
        "GitHub 저장소의 Settings → Secrets and variables → Actions 에 .env 와 동일한 키들을 등록합니다.",
    )
    add_para(doc, "등록할 secrets:", bold=True)
    add_bullet(doc, "GOOGLE_OAUTH_TOKEN (필수)")
    add_bullet(doc, "GEMINI_API_KEY (권장)")
    add_bullet(doc, "GDRIVE_FOLDER_ID (선택, 없으면 categories.yaml 의 folder_id 사용)")
    add_bullet(doc, "GDRIVE_SPREADSHEET_ID (선택)")
    add_bullet(doc, "SUPABASE_URL · SUPABASE_SERVICE_KEY (Supabase 사용 시)")
    add_bullet(doc, "SUPABASE_RETENTION_DAYS (선택, 미설정 시 워크플로가 365 주입)")
    add_para(
        doc,
        "이후 매일 KST 09:07 에 .github/workflows/collect.yml 이 자동 실행됩니다. "
        "수동 실행은 Actions 탭 → “뉴스 자동 수집” → Run workflow.",
    )

    # ── 5. 수집 대상 편집 ─────────────────────────────────
    add_heading(doc, "5. 수집 대상 편집 — config/categories.yaml", 1)
    add_code(
        doc,
        'folder_id: "..."           # 환경변수 없을 때 fallback\n'
        'spreadsheet_id: ""\n'
        "limit_per_source: 6        # 소스당 기본 수집 건수\n\n"
        "categories:\n"
        "  - name: 경제\n"
        "    sources:\n"
        "      - type: rss\n"
        "        url: https://www.hankyung.com/feed/economy\n"
        "        label: 한국경제_경제\n"
        "        crawl_body: true   # 본문까지 크롤",
    )
    add_para(doc, "추가/삭제 시 주의:", bold=True)
    add_bullet(doc, "label 은 시트의 ‘소스’ 컬럼에 그대로 기록 — 중복 안 되게")
    add_bullet(doc, "crawl_body: false 면 RSS description 만 사용 (빠르지만 본문 부실)")
    add_bullet(doc, "카테고리 단위로 limit 을 다르게 지정 가능")
    add_para(
        doc,
        "변경 후 로컬에서 한 번 돌려보고 커밋하는 것을 권장합니다.",
        italic=True,
    )

    # ── 6. 출력 데이터 ────────────────────────────────────
    add_heading(doc, "6. 출력 데이터", 1)

    add_heading(doc, "6-1. Google Sheets", 2)
    add_bullet(doc, "파일명: news-{년도}.xlsx (예: news-2026.xlsx)")
    add_bullet(doc, "탭: 월 단위 — 2026-05, 2026-06, …")
    add_bullet(doc, "컬럼 순서: 수집시간 · 발행시간 · 카테고리 · 제목 · 요약 · 감성 · 중요도 · …")
    add_bullet(
        doc,
        "TubeAI 가 읽을 수 있게 공유 필수: 시트 → 공유 → ‘링크가 있는 모든 사용자(뷰어)’",
    )

    add_heading(doc, "6-2. Supabase", 2)
    add_bullet(doc, "테이블: news_articles (콘솔에서 사전 생성 필요)")
    add_bullet(doc, "url_hash unique key → 중복 자동 스킵")
    add_bullet(doc, "매 실행마다 SUPABASE_RETENTION_DAYS 이전 데이터 자동 cleanup")
    add_bullet(doc, "기본값: 365일 (GitHub Actions) / 90일 (.env 기본)")

    # ── 7. Gemini 쿼터 관리 ──────────────────────────────
    add_heading(doc, "7. Gemini 쿼터 관리", 1)
    add_para(doc, "무료 티어는 분당 5회 / 일일 20회로 타이트합니다. 자동 보호장치:")
    add_table(
        doc,
        ["보호 장치", "동작"],
        [
            ["배치 크기 20", "한 번의 API 호출에 20개 기사를 묶어 처리. 일 400 기사까지 AI 커버."],
            ["Circuit breaker", "429 Quota exceeded 감지 시 남은 기사는 즉시 fallback 으로 전환."],
            ["SKIP_GEMINI=1", "처음부터 AI 호출 스킵, 전부 규칙 기반 fallback."],
        ],
    )
    add_para(
        doc,
        "운영 팁: 일일 한도가 자주 터지면 AI_BATCH_SIZE=30 으로 늘려 호출 횟수를 줄이세요.",
        italic=True,
    )

    # ── 8. 트러블슈팅 ─────────────────────────────────────
    add_heading(doc, "8. 트러블슈팅", 1)

    add_heading(doc, "Actions 가 자동 실행되지 않을 때", 2)
    add_bullet(
        doc,
        "60일 무커밋 → schedule 비활성화: keepalive step 이 매일 LAST_RUN.txt 를 커밋해 방지. 비활성화됐다면 Actions 탭에서 수동 실행으로 재활성화.",
    )
    add_bullet(doc, "기본 브랜치 확인 — on.schedule 은 default branch(main) 의 워크플로만 실행")
    add_bullet(
        doc,
        "invalid_grant 에러 — GOOGLE_OAUTH_TOKEN 의 refresh_token 이 6개월 미사용 또는 revoke. get_token.py 재실행 → secrets 갱신.",
    )
    add_bullet(doc, "시크릿 누락 — Actions 로그 첫 step ‘🔐 시크릿 존재 여부 점검’ 에서 확인")

    add_heading(doc, "429 Quota exceeded 가 자주 보일 때", 2)
    add_bullet(doc, "AI_BATCH_SIZE 를 20 → 30 으로 올려 호출 횟수 감소")
    add_bullet(doc, "AI_CALL_DELAY 를 1.5 → 3 으로 올려 분당 호출 분산")
    add_bullet(doc, "임시: SKIP_GEMINI=1 로 우회하고 fallback 으로 동작")

    add_heading(doc, "Sheets 에 행이 안 쌓일 때", 2)
    add_bullet(doc, "GDRIVE_FOLDER_ID 가 본인 계정에서 접근 가능한 폴더인지")
    add_bullet(doc, "OAuth 스코프에 spreadsheets / drive 둘 다 포함됐는지 (get_token.py 재실행)")
    add_bullet(doc, "로그에 ‘이미 수집됨 (skip)’ 만 보이면 정상 — 어제 수집한 URL 은 중복 차단됨")

    add_heading(doc, "Supabase 에 안 쌓일 때", 2)
    add_bullet(doc, "SUPABASE_URL / SUPABASE_SERVICE_KEY 둘 다 들어가 있어야 활성화")
    add_bullet(doc, "service_role 키여야 함 (anon 키는 RLS 때문에 쓰기 거부됨)")
    add_bullet(doc, "테이블 news_articles 가 사전에 생성돼 있어야 함")

    # ── 9. 폴더 구조 ─────────────────────────────────────
    add_heading(doc, "9. 폴더 구조", 1)
    add_code(
        doc,
        "news-collector/\n"
        "├── .env.example              ← 환경변수 템플릿\n"
        "├── .gitignore\n"
        "├── .github/\n"
        "│   ├── workflows/collect.yml ← GitHub Actions 스케줄\n"
        "│   └── LAST_RUN.txt          ← keepalive 타임스탬프 (자동 커밋)\n"
        "├── config/\n"
        "│   └── categories.yaml       ← RSS 소스 · 카테고리 정의\n"
        "├── src/\n"
        "│   ├── collector.py          ← 수집기 본체\n"
        "│   └── supabase_client.py    ← Supabase 저장·cleanup 로직\n"
        "├── scripts/\n"
        "│   └── backfill_to_supabase.py  ← 시트 → Supabase 1회 백필\n"
        "├── docs/                     ← (구) GitHub Pages 자산 — v2 부터 미사용\n"
        "├── doc/\n"
        "│   └── USAGE.docx            ← 이 문서\n"
        "├── get_token.py              ← OAuth 토큰 1회 발급\n"
        "├── requirements.txt\n"
        "└── README.md                 ← 프로젝트 개요",
    )

    # ── 10. 일상 운영 체크리스트 ──────────────────────────
    add_heading(doc, "10. 일상 운영 체크리스트", 1)
    add_table(
        doc,
        ["빈도", "할 일"],
        [
            ["매일", "(선택) Actions 탭에서 마지막 실행 상태 초록인지 확인"],
            ["주 1회", "Sheets 에서 당일 행이 정상적으로 추가됐는지"],
            ["월 1회", "새 월 탭이 자동 생성됐는지, 헤더 정상인지"],
            ["분기 1회", "Drive 폴더에 새 연단위 파일이 필요한지 (연 바뀌면 자동 생성)"],
            ["6개월 1회", "OAuth refresh_token 만료 임박 — 한 번 더 실행돼 active 유지되는지"],
        ],
    )

    # ── 11. 참고 ─────────────────────────────────────────
    add_heading(doc, "11. 참고", 1)
    add_bullet(doc, "프로젝트 개요 / 변경 이력: README.md")
    add_bullet(doc, "워크플로 정의: .github/workflows/collect.yml")
    add_bullet(doc, "수집 대상: config/categories.yaml")
    add_bullet(doc, "데이터 소비처: YoutubeProgram/TubeAI_app 의 RSS 뉴스 피드 모듈")

    add_divider(doc)
    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = foot.add_run("© news-collector — 운영 가이드")
    foot_run.italic = True
    foot_run.font.size = Pt(9)
    foot_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(OUT_PATH)
    print(f"✅ Generated: {OUT_PATH}")


if __name__ == "__main__":
    build()
